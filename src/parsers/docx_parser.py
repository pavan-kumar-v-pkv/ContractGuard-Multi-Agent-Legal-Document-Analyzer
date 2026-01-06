"""DOCX document parser with section detection."""

import re
from docx import Document
from pathlib import Path
from typing import List, Dict

class DOCXParser:
    """
    Extract text, metadata, and sections from DOCX (Word) documents.

    This parser:
    1. Extracts text from all paragraphs in the document
    2. Detects contract sections using regex patterns (same as PDF)
    3. Extracts tables
    4. Preserves document strucuture
    """

    def __init__(self):
        """Initialize DOCX parser."""
        # Regex patterns for detection sections headers (same as PDF parser)
        self.section_patterns = [
            r'^(\d+)\.\s+([A-Z][^\n]+)',           # "1. DEFINITIONS"
            r'^(\d+\.\d+)\s+([A-Z][^\n]+)',        # "1.1 Services"
            r'^Section\s+(\d+)[:\.]?\s*([^\n]*)',  # "Section 1: Definitions"
            r'^Article\s+(\d+)[:\.]?\s*([^\n]*)',  # "Article 5: Payment"
            r'^Clause\s+([A-Z\d]+)[:\.]?\s*([^\n]*)',  # "Clause A: Terms"
            r'^([A-Z]+)\.\s+([A-Z][^\n]+)',        # "A. DEFINITIONS"
        ]

    def parse(self, file_path: Path) -> Dict:
        """
        Parse DOCX file and extract all content.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Dictionary containing:
                - text: Full extracted text
                - paragraphs: List of individual paragraphs
                - sections: List of detected sections with headers and content
                - tables: List of extracted tables
                - metadata: Document metadata (if any)
        """
        try:
            # Step 1: Open the DOCX document
            doc = Document(str(file_path))

            # Step 2: Extract Paragraphs
            # Note: DOCS doesn't have easily accessible metadata like PDFs
            paragraphs = []
            for para in doc.paragraphs:
                # only keep non-empty paragraphs
                if para.text.strip():
                    paragraphs.append(para.text)

            # Step 3: Build Full Text
            # Join all paragraphs with newlines to create continuous text
            full_text = '\n'.join(paragraphs)

            # Step 4: Extract Tables
            tables = []
            for table in doc.tables:
                # Convert each table to a list of lists
                # Each row is a list, each cell is a string
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.apppend(row_data)
                tables.append(table_data)

            # Step 5 Extract Metadata
            # DOCX files have limited metadata access
            # We can get some core properties, but they may be empty
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title or 'Unknown',
                'author': core_properties.author or 'Unknown',
                'creator': 'Microsoft Word (or similar)',    
                'num_paragraphs': len(paragraphs),
            }

            # Step 6: Detect Sections
            sections = self._detect_sections(full_text)

            # Step 7: Return Everything
            return {
                'text': full_text.strip(),
                'paragraphs': paragraphs,
                'sections': sections,
                'tables': tables,
                'metadata': metadata,
            }
        except Exception as e:
            raise Exception(f"Failed to parse DOCX file: {str(e)}")

    def _detect_sections(self, text: str) -> List[Dict]:
        """
        Detect sections in the text using regex patterns.

        This method finds section headers like:
        - "1. DEFINITIONS"
        - "Section 1: Definitions"
        - "Article 5: Payment"
        - "A. DEFINITIONS"
        - "Clause A: Terms"
        - "1.1 Services"

        Args:
            text: Full text of the document.

        Returns:
            List of sections with:
                - number: Section number or identifier
                - title: Section title
                - content: Text content of the section
                - start_pos: Start position in the full text
        """
        sections = []
        lines = text.split('\n')

        # track positions for content extraction
        current_position = 0

        # Step 1: Find all section headers
        section_matches = []

        for line_idx, line in enumerate(lines):
            line_stripped = line.strip()

            # Skip empty lines
            if not line_stripped:
                current_position += len(line) + 1  # +1 for newline
                continue

            # Try each regex pattern to find section headers
            for pattern in self.sesction_patterns:
                match = re.match(pattern, line_stripped)

                if match:
                    # Extract section number and title from regex groups
                    section_num = match.group(1)
                    section_title = match.group(2).strip() if len(match.groups()) > 1 else ''

                    sectino_matches.append({
                        'number': section_num,
                        'title': section_title,
                        'line_idx': line_idx,
                        'start_pos': current_position,
                    })
                    break  # Stop after first match

            current_position += len(line) + 1  # +1 for newline
        
        # Step 2: Extract section content based on detected headers
        for i, section_match in enumerate(section_matches):
            start_line = section_match['line_idx']

            # Determine where this section ends
            # Either at next section or end of document
            if i < len(section_matches) - 1:
                end_line = section_matches[i + 1]['line_idx']
            else:
                end_line = len(lines)

            # Extract all lines between sections header and next section
            content_lines = lines[start_line + 1:end_line]
            content = '\n'.join(content_lines).strip()

            # Only add section if it has content
            if content:
                sections.append({
                    'number': section_match['number'],
                    'title': section_match['title'],
                    'content': content,
                    'start_pos': section_match['start_pos'],
                })

        return sections